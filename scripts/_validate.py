# -*- coding: utf-8 -*-
import sys, re
from docx import Document

files = {
    "MetaProfile软件用户手册_0616v2.docx": "manual",
    "系统测试说明_0616v2.docx": "plan",
    "系统测试报告_0616v2.docx": "report",
}
for f, tag in files.items():
    print(f"\n===== {tag}: {f} =====")
    d = Document(f)
    n_img = len(d.inline_shapes)
    n_tbl = len(d.tables)
    heads = []
    bad_num = []
    for p in d.paragraphs:
        nm = p.style.name if p.style else ""
        if re.match(r"^Heading \d+$", nm):
            t = p.text.strip()
            heads.append((nm, t[:40]))
            if re.match(r"^\d+(\.\d+)*\s+\S", t):
                bad_num.append(t[:40])
    print(f"  inline_images={n_img}  tables={n_tbl}  headings={len(heads)}")
    print(f"  first 8 headings:")
    for nm, t in heads[:8]:
        print(f"    [{nm}] {t}")
    print(f"  headings-with-manual-number(应为0): {len(bad_num)}")
    for t in bad_num[:5]:
        print(f"    !! {t}")
