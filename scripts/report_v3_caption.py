# -*- coding: utf-8 -*-
"""
编辑《系统测试报告_0616v3.docx》图标题：
  1) 全部图标题（文本以「图 」开头的段落）按文档顺序重编号为 图1..图N；
     旧「图 1/2/3」数字编号被替换；保留「TC-XXX」等标识作描述。
  2) 图标题段落居中对齐。
就地保存。
"""
from __future__ import annotations
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = "系统测试报告_0616v3.docx"
CAP_RE = re.compile(r"^图\s+(.*)$")
NUM_TOKEN_RE = re.compile(r"^\d+\s*(.*)$", re.S)


def set_para_text(p, text: str) -> None:
    """保留首 run 格式，整段文本替换为 text。"""
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(text)


def main() -> int:
    doc = Document(PATH)
    n = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        m = CAP_RE.match(t)
        if not m:
            continue
        rest = m.group(1).strip()
        mn = NUM_TOKEN_RE.match(rest)        # 去掉旧的纯数字编号（图 1/2/3）
        desc = mn.group(1).strip() if mn else rest
        n += 1
        newtext = f"图{n}  {desc}".rstrip()
        set_para_text(p, newtext)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(PATH)
    print(f"captions renumbered+centered: {n}")
    return 0


if __name__ == "__main__":
    sys.exit(main()) if False else main()
