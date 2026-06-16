# -*- coding: utf-8 -*-
"""
v3 = 克隆《系统测试报告模板.docx》（GJB 模板壳：封面/签发页/文档修改记录/TOC/样式）
     + 填入 v1 全量内容（修订记录/两轮汇总/1~6 章 R1 + R2，含 72 截图 + 75 表）。

- 标题样式映射：v1 H1/H2/H3 -> 模板 id 2/4/5（均自动编号）。
- v1 标题文本含手动编号（'1 测试概述'），剥掉前导 'N ' / 'N.M ' 防双编号。
- 正文 -> 模板 'Normal Indent'(id 3)。
- 表 -> 重建为模板 'Table Grid'(id 23)。
- 图片：取 v1 image part blob + extent，按比例插入（竖图限高 16cm，横图限宽 15cm）。
"""
from __future__ import annotations
import io, re, shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

TPL = Path("系统测试报告模板.docx")
SRC = Path("系统测试报告_0616v1.docx")
OUT = Path("系统测试报告_0616v3.docx")

# v1 style name -> template style id
STYLE_MAP = {"Heading 1": "2", "Heading 2": "4", "Heading 3": "5", "Heading 4": "6"}
NUM_RE = re.compile(r"^\d+(\.\d+)*\s+")


def set_pstyle(p, sid: str) -> None:
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:pStyle")):
        pPr.remove(e)
    ps = OxmlElement("w:pStyle"); ps.set(qn("w:val"), sid); pPr.insert(0, ps)


def clear_body(doc) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    for idx, ch in enumerate(children):
        if ch.tag.split("}")[-1] != "p":
            continue
        para = next((p for p in doc.paragraphs if p._p is ch), None)
        nm = para.style.name if para and para.style else ""
        if re.match(r"^Heading \d+$", nm):
            start = idx; break
    if start is None:
        raise RuntimeError("模板未找到正文起始 Heading")
    for ch in children[start:]:
        tag = ch.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        if tag == "p":
            pPr = ch.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue
        body.remove(ch)


def extract_images(v1doc, p):
    """返回该段内 (blob_bytes, width_emu, height_emu) 列表。"""
    out = []
    for d in p._p.findall(".//" + qn("w:drawing")):
        ext = d.find(".//" + qn("wp:extent"))
        cx = int(ext.get("cx")) if ext is not None else 0
        cy = int(ext.get("cy")) if ext is not None else 0
        blip = d.find(".//" + qn("a:blip"))
        rid = blip.get(qn("r:embed")) if blip is not None else None
        if rid and rid in v1doc.part.rels:
            part = v1doc.part.related_parts[rid]
            out.append((part.blob, cx, cy))
    return out


def add_table(tdoc, v1tbl):
    rows = [[c.text for c in row.cells] for row in v1tbl.rows]
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = tdoc.add_table(rows=len(rows), cols=ncols)
    t.style = tdoc.styles["Table Grid"]
    t.autofit = True
    for i, row in enumerate(rows):
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            cell = t.rows[i].cells[j]
            cell.text = val


def main() -> None:
    tdoc = Document(str(TPL))
    clear_body(tdoc)
    v1 = Document(str(SRC))

    src_children = list(v1.element.body.iterchildren())
    # 找到 v1 正文起点：第一个 Heading（'修订记录'），跳过其封面 Title/Normal 行
    start = None
    for idx, ch in enumerate(src_children):
        if ch.tag.split("}")[-1] != "p":
            continue
        para = next((p for p in v1.paragraphs if p._p is ch), None)
        if para and para.style and re.match(r"^Heading \d+$", para.style.name):
            start = idx; break
    if start is None:
        raise RuntimeError("v1 未找到正文起始 Heading")

    for ch in src_children[start:]:
        tag = ch.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        if tag == "tbl":
            # 定位对应 Table 对象
            tbl = next((t for t in v1.tables if t._tbl is ch), None)
            if tbl is not None:
                add_table(tdoc, tbl)
            continue
        if tag != "p":
            continue
        para = next((p for p in v1.paragraphs if p._p is ch), None)
        if para is None:
            continue
        name = para.style.name if para.style else ""
        text = para.text
        # 图片段
        imgs = extract_images(v1, para)
        if imgs:
            for blob, cx, cy in imgs:
                ip = tdoc.add_paragraph()
                ip.alignment = 1
                run = ip.add_run()
                if cy and cx:
                    w_cm = cx / 360000.0; h_cm = cy / 360000.0
                    if h_cm > w_cm:  # 竖图
                        run.add_picture(io.BytesIO(blob), height=Cm(min(h_cm, 16)))
                    else:
                        run.add_picture(io.BytesIO(blob), width=Cm(min(w_cm, 15)))
                else:
                    run.add_picture(io.BytesIO(blob), width=Cm(14))
            # 若图片段还有说明文字，跟在图后
            if text.strip():
                bp = tdoc.add_paragraph(); set_pstyle(bp, "3"); bp.add_run(text.strip())
            continue
        # 标题
        if name in STYLE_MAP:
            clean = NUM_RE.sub("", text.strip())
            hp = tdoc.add_paragraph(); set_pstyle(hp, STYLE_MAP[name]); hp.add_run(clean)
            continue
        # 普通正文（Normal/Title/其它）-> Normal Indent
        if not text.strip():
            continue
        bp = tdoc.add_paragraph(); set_pstyle(bp, "3"); bp.add_run(text.strip())

    # 让 Word 重建目录
    settings = tdoc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true"); settings.append(uf)

    tdoc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
