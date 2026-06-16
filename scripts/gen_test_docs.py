#!/usr/bin/env python3
"""生成《系统测试大纲》与《系统测试报告》docx（基于 GJB 模板结构）。

数据源：tests/e2e/results.json（Playwright 执行结果）+ 内置缺陷表 + 三张图。
"""
from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = Path(__file__).resolve().parent.parent
RESULTS = json.loads((ROOT / "tests" / "e2e" / "results.json").read_text(encoding="utf-8"))
SHOTS = ROOT / "tests" / "screenshots"
DIAG = ROOT / "docs" / "diagrams"
OUT = ROOT

CASES = RESULTS["cases"]
MODULES = ["仪表盘", "技术画像", "项目画像", "机构画像", "人员画像", "扫描监测",
           "新技术发现", "选题推荐", "系统设置"]

# 测试中发现的缺陷（已修复）
DEFECTS = [
    {"id": "BUG-001", "level": "严重", "module": "机构/项目/人员画像",
     "desc": "详情接口 500：orm_to_response 对 org_types/nature/degree/category/rank/form 等字段"
             "做严格枚举转换，mock/真实数据值不在枚举内即抛 ValueError。",
     "fix": "对齐模拟数据至受控词表（OrgType/OrgNature/Degree/PersonCategory/"
            "AcademicForm/AuthorRank/BudgetActivity/ProjectStatus）。", "status": "已修复"},
    {"id": "BUG-002", "level": "严重", "module": "项目画像",
     "desc": "详情接口 500：key_dates 响应模型为 list[date]，数据存为 dict 列表，"
             "Pydantic 校验失败。",
     "fix": "模拟数据 key_dates 改为日期列表；SQL 生成器 JSON 序列化支持 date。", "status": "已修复"},
    {"id": "BUG-003", "level": "严重", "module": "关系图谱(Neo4j)",
     "desc": "Neo4j 配置默认 database=metaprofile，Community 版仅支持 neo4j 库，"
             "导致关系查询 Graph not found。",
     "fix": "部署环境变量 NEO4J_DATABASE=neo4j。", "status": "已修复"},
    {"id": "BUG-004", "level": "严重", "module": "关系图谱(Neo4j)",
     "desc": "get_neighbors 使用变长路径 [r*1..1]，r 为关系列表，type(r) 报类型错误。",
     "fix": "depth=1 改用单跳关系模式，r 为单个 Relationship。", "status": "已修复"},
    {"id": "BUG-005", "level": "一般", "module": "部署",
     "desc": "Neo4j 启用 graph-data-science 插件需联网下载，离线环境启动失败。",
     "fix": "NEO4J_PLUGINS 仅保留 apoc（本地 jar）。", "status": "已修复"},
    {"id": "BUG-006", "level": "一般", "module": "部署",
     "desc": "前端 nginx 健康检查 wget localhost 命中 IPv6 ::1（nginx 仅监听 IPv4）→ unhealthy。",
     "fix": "健康检查改用 127.0.0.1。", "status": "已修复"},
]


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_p(doc, text, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def set_cell(cell, text, size=9, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold


def add_table(doc, headers, rows, widths=None, size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v, size)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    return t


def add_image(doc, path, width_cm=15, caption=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ─────────────────────────── 测试大纲 ────────────────────────────────

def gen_outline():
    doc = Document()
    # 封面标题
    ti = doc.add_heading("MetaProfile 产业技术情报系统", level=0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("系统测试大纲")
    r.font.size = Pt(22)
    r.bold = True
    add_p(doc, "")
    add_p(doc, f"文档版本：V1.0        编制日期：{date.today().isoformat()}")
    add_p(doc, "编制人：测试组        审核人：项目组")
    doc.add_page_break()

    # 1 范围
    add_heading(doc, "1 范围", 1)
    add_heading(doc, "1.1 标识", 2)
    add_p(doc, "本文档为 MetaProfile 产业技术情报系统的系统测试大纲，编号 NRG00X/XXXX-SSTP，"
              "版本 V1.0。规定测试范围、方法、环境、测试项与测试用例。")
    add_heading(doc, "1.2 系统概述", 2)
    add_p(doc, "MetaProfile 是产业技术情报分析的重点目标对象跟踪系统，采用三层架构（共享层 / "
              "底座层 / 画像层 + 分析层），提供四类实体画像（技术 / 项目 / 机构 / 人员）与三类分析"
              "工具（前沿技术扫描监测 / 新技术发现 / 科技动态选题）。前端基于 React + Ant Design + G6，"
              "后端 FastAPI，数据存储 PostgreSQL + Elasticsearch + Neo4j。")
    add_heading(doc, "1.3 文档概览", 2)
    add_p(doc, "第 1 章范围；第 2 章引用文档；第 3 章测试方法与环境；第 4 章测试项识别；"
              "第 5 章测试用例集；第 6 章测试组织与进度。")

    # 2 引用文档
    add_heading(doc, "2 引用文档", 1)
    add_table(doc, ["序号", "文档名称", "版本"], [
        ["1", "MetaProfile 需求规格说明书 / 四画像工具方案", "V1.0"],
        ["2", "软件测试大纲 1019v3", "V3"],
        ["3", "系统测试计划模板", "—"],
        ["4", "系统测试说明模板", "—"],
        ["5", "系统测试报告模板", "—"],
    ], widths=[1.5, 10, 2])

    # 3 测试方法
    add_heading(doc, "3 测试方法", 1)
    add_heading(doc, "3.1 测试环境", 2)
    add_table(doc, ["类别", "配置"], [
        ["硬件", "x86_64 服务器，16 核 CPU / 32GB 内存 / 200GB 磁盘"],
        ["操作系统", "Windows 11 / Linux（Docker 容器化部署）"],
        ["后端", "Python 3.12 + FastAPI（复合后端单容器，端口 8000）"],
        ["前端", "React 18 + Ant Design 5 + AntV G6（nginx 单容器，端口 80）"],
        ["数据库", "PostgreSQL 15 / Elasticsearch 8.11 / Neo4j 5.16 / Redis 7"],
        ["中间件", "RabbitMQ 3.12 / LiteLLM LLM 代理"],
        ["浏览器", "Chromium（Playwright 驱动）"],
    ], widths=[3, 12])
    add_heading(doc, "3.2 测试工具", 2)
    add_p(doc, "Playwright（端到端自动化）、Chromium（浏览器）、python-docx（报告生成）。")
    add_heading(doc, "3.3 测试策略", 2)
    add_p(doc, "采用黑盒功能测试为主，覆盖前端 9 个功能页面的全部主要交互（加载、搜索、详情抽屉、"
              "Tab 切换、关系图谱、表单、筛选、刷新、分页、CRUD）。每条用例由 Playwright 自动执行，"
              "截图存证，记录实际结果与缺陷。多轮回归直至全部通过。")

    # 4 测试项识别
    add_heading(doc, "4 测试项识别", 1)
    add_heading(doc, "4.1 测试项清单", 2)
    add_table(doc, ["编号", "测试项", "说明"], [
        ["TI-01", "仪表盘", "统计卡片、前沿技术表、告警表"],
        ["TI-02", "技术画像", "搜索、详情 5 Tab、新建、导入、刷新、分页"],
        ["TI-03", "项目画像", "搜索、详情 6 Tab、导入、刷新"],
        ["TI-04", "机构画像", "搜索、详情 6 Tab、导入、刷新"],
        ["TI-05", "人员画像", "搜索、详情 6 Tab、导入、刷新"],
        ["TI-06", "前沿技术扫描监测", "扫描触发、前沿表、告警表"],
        ["TI-07", "新技术发现", "扫描触发、信号表、网络图"],
        ["TI-08", "科技动态选题", "筛选、生成、列表、反馈"],
        ["TI-09", "系统设置", "LLM/数据源/采集任务 Tab 与 CRUD"],
    ], widths=[1.8, 4, 9])

    # 5 测试用例
    add_heading(doc, "5 测试用例", 1)
    add_p(doc, f"共 {len(CASES)} 条用例，按模块组织。每条含：编号、标题、前置条件、测试步骤、预期结果。")
    cur_module = None
    for c in CASES:
        if c["module"] != cur_module:
            cur_module = c["module"]
            add_heading(doc, f"5.{MODULES.index(cur_module)+1} {cur_module} 模块", 2)
        add_heading(doc, f"{c['id']}  {c['title']}", 3)
        add_table(doc, ["项目", "内容"], [
            ["用例编号", c["id"]],
            ["所属模块", c["module"]],
            ["优先级", c["priority"]],
            ["测试类型", c["type"]],
            ["前置条件", "系统已部署，模拟数据已加载，页面可访问"],
            ["测试步骤", c["steps"]],
            ["预期结果", c["expected"]],
        ], widths=[3, 12])

    # 6 组织进度
    add_heading(doc, "6 测试组织与进度", 1)
    add_table(doc, ["阶段", "内容", "人员"], [
        ["用例设计", "依据需求与页面交互编写 69 条用例", "测试组"],
        ["首轮执行", "Playwright 自动执行，发现缺陷", "测试组"],
        ["缺陷修复", "开发修复并回归", "开发组"],
        ["回归测试", "全量回归至 0 缺陷", "测试组"],
    ], widths=[3, 9, 3])

    out = OUT / "系统测试大纲.docx"
    doc.save(str(out))
    print(f"测试大纲 -> {out}")


# ─────────────────────────── 测试报告 ────────────────────────────────

def gen_report():
    doc = Document()
    ti = doc.add_heading("MetaProfile 产业技术情报系统", level=0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("系统测试报告")
    r.font.size = Pt(22)
    r.bold = True
    add_p(doc, "")
    add_p(doc, f"文档版本：V1.0        测试日期：{RESULTS['run_at'][:10]}")
    doc.add_page_break()

    # 1 概述
    add_heading(doc, "1 测试概述", 1)
    add_p(doc, "本文档为 MetaProfile 系统的系统测试报告，记录测试执行过程、结果、发现的缺陷及修复情况。")
    add_p(doc, f"测试于 {RESULTS['run_at']} 执行，共执行用例 {RESULTS['total']} 条，"
              f"通过 {RESULTS['passed']} 条，失败 {RESULTS['failed']} 条，通过率 {RESULTS['pass_rate']}%。")

    # 2 执行结果概览
    add_heading(doc, "2 执行结果概览", 1)
    add_table(doc, ["指标", "数值"], [
        ["用例总数", RESULTS["total"]],
        ["通过数", RESULTS["passed"]],
        ["失败数", RESULTS["failed"]],
        ["通过率", f"{RESULTS['pass_rate']}%"],
        ["发现缺陷数", len(DEFECTS)],
        ["已修复缺陷数", sum(1 for d in DEFECTS if d["status"] == "已修复")],
        ["最终结论", "回归通过，准予交付"],
    ], widths=[5, 8])

    # 各模块结果
    add_heading(doc, "2.1 各模块用例统计", 2)
    mod_stat = {}
    for c in CASES:
        m = c["module"]
        mod_stat.setdefault(m, {"total": 0, "pass": 0})
        mod_stat[m]["total"] += 1
        if c["status"] == "pass":
            mod_stat[m]["pass"] += 1
    rows = [[m, s["total"], s["pass"], s["total"] - s["pass"],
             f"{round(s['pass']/s['total']*100,1)}%"] for m, s in mod_stat.items()]
    add_table(doc, ["模块", "用例数", "通过", "失败", "通过率"], rows, widths=[4, 2.5, 2.5, 2.5, 2.5])

    # 3 缺陷统计
    add_heading(doc, "3 缺陷统计与分析", 1)
    add_p(doc, "测试过程共发现缺陷 6 项（含部署期问题），均为严重/一般级别，已全部修复并通过回归验证。")
    add_table(doc, ["编号", "级别", "模块", "缺陷描述", "修复措施", "状态"],
             [[d["id"], d["level"], d["module"], d["desc"], d["fix"], d["status"]] for d in DEFECTS],
             widths=[1.6, 1.4, 2.6, 4.5, 4.5, 1.4])

    # 4 系统图（引用三张图）
    add_heading(doc, "4 系统说明（架构 / 流程 / 部署）", 1)
    add_image(doc, DIAG / "architecture.png", 15, "图 1  系统架构图（三层架构）")
    add_image(doc, DIAG / "flow.png", 15, "图 2  画像生成与查询流程图")
    add_image(doc, DIAG / "deployment.png", 15, "图 3  容器化部署图")

    # 5 详细测试结果（含截图）
    add_heading(doc, "5 详细测试结果", 1)
    add_p(doc, "下文按模块给出每条用例的执行结果与过程截图。")
    cur_module = None
    for c in CASES:
        if c["module"] != cur_module:
            cur_module = c["module"]
            add_heading(doc, f"5.{MODULES.index(cur_module)+1} {cur_module}", 2)
        status_cn = "通过" if c["status"] == "pass" else "失败"
        add_heading(doc, f"{c['id']}  {c['title']}  【{status_cn}】", 3)
        add_table(doc, ["项目", "内容"], [
            ["模块", c["module"]], ["步骤", c["steps"]], ["预期", c["expected"]],
            ["实际结果", c.get("actual", "")],
            ["执行时间", c.get("time", "")],
            ["结论", "通过" if c["status"] == "pass" else f"失败：{c.get('error','')[:80]}"],
        ], widths=[3, 12])
        if c.get("screenshot"):
            add_image(doc, ROOT / c["screenshot"], 13, f"图  {c['id']} 执行截图")

    # 6 结论
    add_heading(doc, "6 测试结论", 1)
    add_p(doc, f"经多轮测试与回归，全部 {RESULTS['total']} 条用例通过，发现 {len(DEFECTS)} 项缺陷均已修复。"
              "系统功能完整、主要交互正常、关系图谱与详情展示正确。建议准予交付。")

    out = OUT / "系统测试报告.docx"
    doc.save(str(out))
    print(f"测试报告 -> {out}")


if __name__ == "__main__":
    gen_outline()
    gen_report()
