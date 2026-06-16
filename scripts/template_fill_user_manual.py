"""
基于参考模板《软件用户手册.docx》生成《MetaProfile软件用户手册.docx》。

保真策略：克隆参考 docx（保留封面、自动目录 TOC 域、页眉页脚、页码、
全部样式定义、页面设置），仅删除"目录之后"的正文内容，再用模板的
**精确样式 id**（标题 1-5=1/2/3/4/5，正文=a4）填入新内容；表格补边框；
置 updateFields 让 Word 打开时重建目录。
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REF = Path("软件用户手册.docx")
OUT = Path("MetaProfile软件用户手册.docx")
SRC_MD = Path("docs/MetaProfile软件用户手册.md")

# md 标题层级 → 模板样式 id（章节为 heading1，其余递降）
MD_TO_STYLE = {2: "1", 3: "2", 4: "3", 5: "4"}

CONTENT_HEADING_IDS = {"1", "2", "3", "4", "5"}
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


# ---------- XML helpers ----------
def set_pstyle(p, style_id: str) -> None:
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:pStyle")):
        pPr.remove(e)
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)


def shade_para(p, fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def mono_run(p, text: str, size: int = 9) -> None:
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = __import__("docx").shared.Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rfonts.set(qn("w:eastAsia"), "Consolas")


def add_rich(p, text: str) -> None:
    """行内 **bold** 与 `code`。"""
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        else:
            mono_run(p, tok[1:-1])
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "808080")
        borders.append(e)
    tblPr.append(borders)


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# ---------- content builders ----------
def add_heading(doc, text: str, style_id: str) -> None:
    p = doc.add_paragraph()
    set_pstyle(p, style_id)
    p.add_run(text)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph()
    set_pstyle(p, "a4")
    add_rich(p, text)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph()
    set_pstyle(p, "a4")
    add_rich(p, "• " + text)


def add_code(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_pstyle(p, "a4")
    shade_para(p, "F5F5F5")
    for i, ln in enumerate(lines):
        mono_run(p, ln)
        if i < len(lines) - 1:
            p.add_run().add_break()


def add_md_table(doc, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    table_borders(t)
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        shade_cell(c, "E8EEF7")
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h); r.bold = True
    for row in body:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(cells):
                cells[j].paragraphs[0].clear()
                add_rich(cells[j].paragraphs[0], val)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and set(s.replace("|", "").replace("-", "").replace(":", "").strip()) <= set() or (
        re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", s) and "-" in s
    )


# ---------- main ----------
def clear_content(doc) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    for idx, ch in enumerate(children):
        if ch.tag.split("}")[-1] != "p":
            continue
        pPr = ch.find(qn("w:pPr"))
        if pPr is None:
            continue
        ps = pPr.find(qn("w:pStyle"))
        if ps is not None and ps.get(qn("w:val")) in CONTENT_HEADING_IDS:
            start = idx
            break
    if start is None:
        raise RuntimeError("未找到正文起始标题，模板结构异常")
    for ch in children[start:]:
        tag = ch.tag.split("}")[-1]
        if tag == "sectPr":
            continue  # 保留末尾分节
        if tag == "p":
            pPr = ch.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue  # 保留承载分节的段落
        body.remove(ch)


def fill_from_md(doc, md: Path) -> None:
    lines = md.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    started = False  # 跳过封面已有的 H1 + 元信息表，从首个 ## 章节开始
    while i < n:
        line = lines[i]
        s = line.strip()
        if s.startswith("## ") and not started:
            started = True
        if not started:
            i += 1
            continue

        if s.startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            add_code(doc, block)
            continue
        m = re.match(r"^(#{2,5})\s+(.*)$", s)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), MD_TO_STYLE[level])
            i += 1
            continue
        if s.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                if not is_sep(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_md_table(doc, rows)
            continue
        if s.startswith("- "):
            add_bullet(doc, s[2:]); i += 1; continue
        if s.startswith("> "):
            add_body(doc, s[2:]); i += 1; continue
        if not s:
            i += 1
            continue
        add_body(doc, s)
        i += 1


def set_update_fields(doc) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)


def bump_cover_version(doc) -> None:
    """封面占位版本号 0.0.0.1 → V1.0（best-effort 文本替换）。"""
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text and "0.0.0.1" in r.text:
                r.text = r.text.replace("0.0.0.1", "V1.0")
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text and "0.0.0.1" in r.text:
                            r.text = r.text.replace("0.0.0.1", "V1.0")


def main() -> None:
    doc = Document(str(REF))
    clear_content(doc)
    fill_from_md(doc, SRC_MD)
    bump_cover_version(doc)
    set_update_fields(doc)
    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    sys.exit(main())
