"""
Markdown → DOCX 转换器（python-docx，无 pandoc 依赖）。

仅覆盖本项目技术手册用到的 md 子集：
  # / ## / ### 标题、| 表格 |、``` 代码块、- 列表、> 引用、
  段落、行内 `code` 与 **bold**。

用法：python md_to_docx.py <input.md> <output.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def set_cjk_font(doc: Document) -> None:
    """Normal 样式设中文字体，避免 Word 落回默认导致 CJK 显示异常。"""
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_rich(paragraph, text: str) -> None:
    """把含 `code` 与 **bold** 的行内文本拆成多个 run。"""
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), "Consolas")
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def parse_table_row(line: str) -> list[str]:
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


def is_separator_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", line)) and "-" in line and set(line.replace("|", "").replace(":", "").replace("-", "").strip()) <= set()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(12)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    # 段落浅灰底
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    for i, ln in enumerate(lines):
        run = p.add_run(ln)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "Consolas")
        if i < len(lines) - 1:
            run.add_break()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        shade_cell(cell, "E8EEF7")
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(cell_text)
        run.bold = True
    for row_data in body:
        cells = table.add_row().cells
        for j, cell_text in enumerate(row_data):
            if j < len(cells):
                cells[j].paragraphs[0].clear()
                add_rich(cells[j].paragraphs[0], cell_text)


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    set_cjk_font(doc)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            add_code_block(doc, block)
            continue

        # 标题（长前缀优先，避免 #### 误匹配 ###）
        if stripped.startswith("##### "):
            doc.add_heading(stripped[6:], level=5)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < n and is_separator_row(lines[i + 1]):
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                if not is_separator_row(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # 引用
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 列表项
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, stripped[2:])
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_rich(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"saved: {docx_path}  paragraphs/tables rendered")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: python md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
