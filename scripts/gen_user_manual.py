#!/usr/bin/env python3
"""生成《软件用户手册》docx（基于 软件用户手册.docx 模板结构）。

内容：需求与业务场景、9 个功能页面使用、安装部署升级步骤、测试截图、三张系统图。
"""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "tests" / "screenshots"
DIAG = ROOT / "docs" / "diagrams"
OUT = ROOT / "MetaProfile_软件用户手册.docx"


def add_p(doc, text, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def set_cell(cell, text, size=9.5, bold=False):
    cell.text = ""
    r = cell.paragraphs[0].add_run(str(text))
    r.font.size = Pt(size)
    r.bold = bold


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Cm(w)
    return t


def add_image(doc, path, width_cm=14.5, caption=None):
    if Path(path).exists():
        doc.add_picture(str(path), width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(caption)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def shot(name):
    return SHOTS / f"{name}.png"


def build():
    doc = Document()
    ti = doc.add_heading("MetaProfile 产业技术情报系统", level=0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件用户手册")
    r.font.size = Pt(22)
    r.bold = True
    add_p(doc, "")
    add_p(doc, f"文档版本：V1.0        编制日期：{date.today().isoformat()}")
    add_p(doc, "适用对象：情报分析师 / 研究人员 / 系统管理员")
    doc.add_page_break()

    # ── 1 范围 ──
    doc.add_heading("1 范围", 1)
    doc.add_heading("1.1 标识", 2)
    add_p(doc, "本文档为 MetaProfile 产业技术情报系统的软件用户手册，版本 V1.0。")
    doc.add_heading("1.2 系统概述", 2)
    add_p(doc, "MetaProfile 是面向产业技术情报分析的“重点目标对象跟踪系统”，提供四类实体画像"
              "（技术 / 项目 / 机构 / 人员）的构建、检索、关系图谱展示，以及三类分析能力"
              "（前沿技术扫描监测 / 新技术发现 / 科技动态选题）。系统采用前后端分离 + 容器化部署，"
              "支持数据采集、清洗、实体抽取、画像生成与情报分析的端到端流程。")
    doc.add_heading("1.3 文档概述", 2)
    add_p(doc, "第 1 章范围；第 2 章软件综述；第 3 章安装部署与升级；第 4 章软件使用指南"
              "（含各功能页面操作与典型业务场景）。")

    # ── 2 软件综述 ──
    doc.add_heading("2 软件综述", 1)
    doc.add_heading("2.1 软件应用", 2)
    add_p(doc, "用于跟踪重点技术 / 项目 / 机构 / 人员的动态，构建关系图谱，识别前沿技术与弱信号，"
              "辅助情报分析与选题决策。")
    doc.add_heading("2.2 软件清单", 2)
    add_table(doc, ["组件", "技术", "说明"], [
        ["前端", "React 18 + Ant Design 5 + AntV G6", "单页面应用，nginx 托管"],
        ["后端", "Python 3.12 + FastAPI", "复合后端单镜像，77 个 REST 接口"],
        ["数据库", "PostgreSQL 15", "结构化属性存储"],
        ["检索", "Elasticsearch 8.11", "全文 + 向量检索"],
        ["图数据库", "Neo4j 5.16", "关系图谱"],
        ["缓存/队列", "Redis 7 / RabbitMQ 3.12", "缓存与异步任务"],
        ["LLM", "LiteLLM 代理 → Qwen2.5-72B / BGE", "抽取与补全"],
    ], widths=[3, 6, 6])
    doc.add_heading("2.3 软件环境", 2)
    add_table(doc, ["类别", "要求"], [
        ["硬件", "8 核以上 CPU / 16GB 以上内存 / 100GB 磁盘"],
        ["操作系统", "Linux / Windows（Docker 环境）"],
        ["依赖软件", "Docker 24+ / Docker Compose v2"],
        ["浏览器", "Chrome / Edge / Firefox 最新版"],
    ], widths=[3, 12])
    doc.add_heading("2.4 系统架构", 2)
    add_image(doc, DIAG / "architecture.png", 15, "图 1  系统架构图（三层架构）")

    # ── 3 软件入门（安装部署升级） ──
    doc.add_heading("3 软件入门（安装 / 部署 / 升级）", 1)
    doc.add_heading("3.1 安装前准备", 2)
    add_p(doc, "1. 确认已安装 Docker 与 Docker Compose v2。")
    add_p(doc, "2. 准备部署目录，包含 docker-compose.yml、.env、pg_init、litellm.yaml。")
    add_p(doc, "3. 确认端口 80 / 5432 / 7687 / 9200 等未被占用。")
    doc.add_heading("3.2 部署架构", 2)
    add_image(doc, DIAG / "deployment.png", 15, "图 2  容器化部署图")
    add_p(doc, "系统采用 基础设施（不变）+ 单一后端镜像 + 单一前端镜像 的三层部署：")
    add_table(doc, ["层级", "镜像 / 容器", "说明"], [
        ["基础设施", "postgres / elasticsearch / redis / neo4j / rabbitmq / litellm", "不变"],
        ["后端", "metaprofile-backend（:8000）", "复合 app，承载全部后端功能"],
        ["前端", "metaprofile-frontend（:80）", "nginx 托管 + 反向代理"],
        ["一次性任务", "migrate / seed", "建表 + 灌入模拟数据（复用后端镜像）"],
    ], widths=[2.5, 7, 5])
    doc.add_heading("3.3 安装步骤", 2)
    add_p(doc, "1. 进入 deploy 目录：cd deploy")
    add_p(doc, "2. 构建镜像并启动：docker compose up -d --build")
    add_p(doc, "3. 等待 migrate（建表）与 seed（灌入 100×4 模拟数据 + 关系图谱）自动完成。")
    add_p(doc, "4. 访问 http://localhost 即可使用系统。")
    doc.add_heading("3.4 安装后验证", 2)
    add_p(doc, "1. docker compose ps 确认全部容器 healthy。")
    add_p(doc, "2. 浏览器访问 http://localhost，应看到仪表盘页面与统计数据。")
    add_p(doc, "3. 访问 http://localhost:8000/docs 查看后端 API 文档。")
    doc.add_heading("3.5 启动与停止", 2)
    add_p(doc, "启动：docker compose up -d        停止：docker compose down")
    add_p(doc, "查看日志：docker compose logs -f backend")
    doc.add_heading("3.6 升级步骤", 2)
    add_p(doc, "1. 拉取 / 更新源代码。")
    add_p(doc, "2. docker compose build（重建镜像，基础层缓存复用）。")
    add_p(doc, "3. docker compose up -d（滚动重建变更的容器）。")
    add_p(doc, "4. 若涉及表结构变更，migrate 任务自动执行 alembic 迁移。")
    add_p(doc, "5. 若需重置模拟数据：docker compose run --rm seed。")
    doc.add_heading("3.7 卸载", 2)
    add_p(doc, "docker compose down -v（-v 同时删除数据卷，谨慎使用）。")

    # ── 4 软件使用指南 ──
    doc.add_heading("4 软件使用指南", 1)
    doc.add_heading("4.1 业务流程", 2)
    add_image(doc, DIAG / "flow.png", 15, "图 3  画像生成与查询流程图")
    add_p(doc, "原始资料（论文 / 专利 / 项目 / 机构 / 人员资料）经采集 → 清洗 → 实体识别 → 属性抽取 → "
              "关系抽取 → 消歧融合 → 三库入库 → 画像生成，形成四类画像与关系图谱，供检索、统计、"
              "扫描监测、新技术发现与选题使用。")
    doc.add_heading("4.2 功能页面与操作", 2)

    pages = [
        ("4.2.1 仪表盘", "TC-DASH-01", "登录后首页，展示技术/项目/机构/人员四类画像总量统计卡片，"
         "前沿技术 Top5 表（按融合评分）与最新告警表。用于快速掌握全局态势。"),
        ("4.2.2 技术画像", "TC-TECH-01", "列表页支持关键词搜索（如'芯片'）；点击'详情'打开抽屉，"
         "含 5 个 Tab：基本信息、里程碑（时间线）、科研成果（经费/学术/实验表格）、统计图表、"
         "关联图谱（G6 径向图）。可新建、批量导入、刷新。"),
        ("4.2.3 项目画像", "TC-PROJ-01", "项目列表与搜索；详情抽屉含 6 Tab：基本信息、研究内容、"
         "发展历程、预算、项目成果、关联图谱。展示项目的主管/承研机构、研究人员、涉及技术。"),
        ("4.2.4 机构画像", "TC-ORG-01", "机构列表与搜索；详情抽屉含 6 Tab：基本信息、发展沿革、"
         "科研队伍（团队/设施）、主要成果、荣誉奖励、关联图谱。"),
        ("4.2.5 人员画像", "TC-PER-01", "人员列表与搜索；详情抽屉含 6 Tab：基本信息、工作经历、"
         "教育经历、学术成果、技术关注、关联图谱。展示人员的任职机构、成果与合作关系。"),
        ("4.2.6 前沿技术扫描监测", "TC-SCAN-01", "选择时间范围与技术领域，点击'开始扫描'触发任务；"
         "查看前沿技术清单（融合评分/TRL/状态）与告警列表。"),
        ("4.2.7 新技术发现", "TC-DISC-01", "触发发现扫描，查看弱信号清单（关键词/强度/新颖度/一致性），"
         "可查看信号关联网络图。"),
        ("4.2.8 科技动态选题", "TC-TOP-01", "按状态筛选选题；点击'生成选题'配置数量与周期生成；"
         "查看选题详情并对选题提交评审反馈（结论/评分/意见）。"),
        ("4.2.9 系统设置", "TC-SET-01", "三个 Tab：大模型配置（新增/测试/同步 LiteLLM）、"
         "数据源配置（新增/编辑/删除/立即采集）、采集任务（查看任务与日志）。"),
    ]
    for title, sc, desc in pages:
        doc.add_heading(title, 3)
        add_p(doc, desc)
        add_image(doc, shot(sc), 14, f"图  {title} 页面")

    # 关系图谱特写
    doc.add_heading("4.2.10 关系图谱", 3)
    add_p(doc, "在任一画像详情的'关联图谱'Tab，以 G6 径向布局展示该实体与其他技术/机构/人员/项目的"
              "关系（隶属、承研、管理、涉及、贡献、合作等），节点按实体类型着色。")
    add_image(doc, shot("TC-TECH-09"), 14, "图  关系图谱示例（技术）")

    doc.add_heading("4.3 典型业务场景", 2)
    add_p(doc, "场景一（前沿技术跟踪）：仪表盘查看态势 → 扫描监测触发扫描 → 前沿技术榜单 → "
              "进入技术画像查看里程碑与关系图谱。")
    add_p(doc, "场景二（机构/人员尽调）：机构画像搜索 → 详情查看科研队伍/成果/奖励 → "
              "关联图谱定位合作机构与人员 → 人员画像查看学术成果与技术关注。")
    add_p(doc, "场景三（选题决策）：新技术发现查看弱信号 → 选题服务生成选题 → 评审反馈。")
    doc.add_heading("4.4 数据备份", 2)
    add_p(doc, "PostgreSQL：pg_dump 备份；Neo4j：neo4j-admin dump；Elasticsearch：snapshot API。"
              "建议定期备份 docker 卷（pgdata / neo4jdata / esdata）。")
    doc.add_heading("4.5 常见问题", 2)
    add_table(doc, ["问题", "处理"], [
        ["页面打不开", "docker compose ps 检查 frontend/backend 是否 healthy"],
        ["详情显示加载失败", "检查后端日志与数据库连接；确认 seed 已执行"],
        ["关系图谱为空", "确认 Neo4j healthy 且 NEO4J_DATABASE=neo4j；重跑 seed"],
        ["搜索无结果", "确认模拟数据已加载；重置搜索条件"],
    ], widths=[5, 9])
    doc.add_heading("4.6 快速引用", 2)
    add_table(doc, ["入口", "地址"], [
        ["系统首页", "http://localhost"],
        ["后端 API 文档", "http://localhost:8000/docs"],
        ["Neo4j 浏览器", "http://localhost:7474"],
        ["RabbitMQ 控制台", "http://localhost:15672"],
    ], widths=[5, 9])

    doc.save(str(OUT))
    print(f"用户手册 -> {OUT}")


if __name__ == "__main__":
    build()
