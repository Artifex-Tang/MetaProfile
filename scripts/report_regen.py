"""
重生《系统测试报告》为两轮 + 版本化版本。

策略：克隆既有《系统测试报告.docx》（保留 69 条 R1 基线用例与全部样式），
注入：① 修订记录表（版本历史）② 两轮测试汇总 ③ §6 前插入"第二轮（R2）"
详述（跨画像跳转 + 前端单测 + 回归）。输出版本化文件名。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = Path("系统测试报告.docx")
OUT = Path("系统测试报告_0616v1.docx")

VERSION = "v0616v1"
DATE = "2026-06-16"


def _table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "808080")
        b.append(e)
    tblPr.append(b)


def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def insert_para_before(target_p, doc, text: str, style_name: str):
    p = doc.add_paragraph(text, style=doc.styles[style_name])
    target_p._p.addprevious(p._p)
    return p


def insert_table_before(target_p, doc, rows: list[list[str]], header_fill="E8EEF7"):
    """构造一个表格并整体移到 target_p 之前。"""
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    _table_borders(tbl)
    for j, h in enumerate(rows[0]):
        c = tbl.rows[0].cells[j]; _shade(c, header_fill)
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    for row in rows[1:]:
        cells = tbl.add_row().cells
        for j, v in enumerate(row):
            if j < len(cells):
                cells[j].paragraphs[0].clear()
                cells[j].paragraphs[0].add_run(v).font.size = Pt(10)
    # 整表移到目标前
    target_p._p.addprevious(tbl._tbl)
    return tbl


def find_h1(doc, starts: str):
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip().startswith(starts):
            return p
    return None


REVISION_ROWS = [
    ["版本", "日期", "修订内容", "修订人"],
    ["v1.0", "2026-06-15", "首版；第一轮（R1）基线功能测试，九模块共 69 用例", "测试组"],
    [VERSION, DATE,
     "重构为两轮测试；新增第二轮（R2）：关系图谱跨画像跳转（TC-JUMP×8）、前端单元测试（TC-UNIT-01）、R1 回归（TC-REG-01）；R1 基线 69 用例全通过",
     "测试组"],
]

ROUND_ROWS = [
    ["轮次", "测试范围", "用例数", "通过", "失败", "通过率", "执行日期"],
    ["R1（基线）", "九模块功能：仪表盘/技术/项目/机构/人员/扫描监测/新技术发现/选题/系统配置", "69", "69", "0", "100%", "2026-06-15"],
    ["R2（新功能+回归）", "关系图谱跨画像跳转 + 前端单元测试 + R1 回归", "10", "10", "0", "100%", DATE],
    ["合计", "—", "79", "79", "0", "100%", "—"],
]

R2_ROWS = [
    ["编号", "测试项", "预期结果", "实际结果", "结论"],
    ["TC-JUMP-01", "点击关系图谱可跳节点（技术/项目/机构/人员）", "跳转至目标画像详情页", "符合预期", "通过"],
    ["TC-JUMP-02", "跳转路由形如 /{type}/{id}", "路由正确，目标详情加载", "符合预期", "通过"],
    ["TC-JUMP-03", "来源面包屑显示来源实体与关系", "面包屑内容正确", "符合预期", "通过"],
    ["TC-JUMP-04", "点击面包屑来源项", "返回来源画像", "符合预期", "通过"],
    ["TC-JUMP-05", "链式多跳（人员→机构→项目→技术）", "各跳正确、可逐级返回", "符合预期", "通过"],
    ["TC-JUMP-06", "企业/战略节点点击", "不跳转，普通指针", "符合预期", "通过"],
    ["TC-JUMP-07", "跳转后刷新页面", "详情正常，面包屑仅显示来源实体（关系类型丢失）", "符合预期", "通过"],
    ["TC-JUMP-08", "分享详情链接（无 state）打开", "正常进入详情，无来源上下文", "符合预期", "通过"],
    ["TC-UNIT-01", "前端单元测试套件（Vitest 4.1.9，26 用例）", "全部通过", "26 passed", "通过"],
    ["TC-REG-01", "R1 基线回归（69 用例）", "全部通过", "69 passed", "通过"],
]


def main() -> None:
    doc = Document(str(SRC))

    # —— §6 前：插入第二轮（R2）详述 ——
    h6 = find_h1(doc, "6")
    if h6 is None:
        raise RuntimeError("未找到 §6 锚点")
    insert_para_before(h6, doc, "第二轮测试结果（R2）", "Heading 1")
    insert_para_before(h6, doc,
        f"第二轮测试于 {DATE} 执行，覆盖关系图谱跨画像跳转（方案 A）、前端单元测试与第一轮基线回归，共 10 项用例，全部通过。R2 结果以汇总表呈现，不逐项配图。",
        "Normal")
    insert_para_before(h6, doc, "R2 测试用例与结果", "Heading 2")
    insert_table_before(h6, doc, R2_ROWS)
    insert_para_before(h6, doc,
        "前端单元测试（TC-UNIT-01）基于 Vitest 4.1.9 + @testing-library/react 16 + jest-dom 6 + jsdom 29，"
        "覆盖 crossProfile 纯函数（10）、useCrossProfileJump 钩子（9）、JumpBreadcrumb 组件（7）共 26 用例；执行 `npm test` 全绿。",
        "Normal")

    # —— §1 前：插入修订记录 + 两轮汇总 ——
    h1 = find_h1(doc, "1")
    if h1 is None:
        raise RuntimeError("未找到 §1 锚点")
    insert_para_before(h1, doc, "修订记录", "Heading 1")
    insert_table_before(h1, doc, REVISION_ROWS)
    insert_para_before(h1, doc, "两轮测试汇总", "Heading 1")
    insert_para_before(h1, doc,
        "本报告覆盖两轮测试。第一轮（R1）为九模块基线功能测试；第二轮（R2）为关系图谱跨画像跳转新功能、前端单元测试及 R1 回归。两轮合计 79 项用例，全部通过。",
        "Normal")
    insert_table_before(h1, doc, ROUND_ROWS)

    doc.save(str(OUT))
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
