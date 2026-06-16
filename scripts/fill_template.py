# -*- coding: utf-8 -*-
"""
通用模板填充器：克隆指定 docx 模板（保封面/签发页/TOC 域/页眉页脚/全部样式），
删除"目录之后首个 Heading 起"的正文，再按 md 内容用模板**精确样式 id** 填充。

标题样式均带自动编号（numPr），故 md 标题文本**不含手动编号**。
图片用 ![caption](relative/path) 表示，居中插入并附图注。

用法：python fill_template.py <manual|plan|report>
配置见 DOCS。
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent

DOCS = {
    "manual": {
        "template": "软件用户手册.docx",
        "md": "docs/MetaProfile软件用户手册.md",
        "out": "MetaProfile软件用户手册_0616v2.docx",
        "head": {2: "1", 3: "2", 4: "3", 5: "4"},   # md H2..H5 -> Heading1..4
        "body": "a4",
        "version_replaces": [("0.0.0.1", "V1.0")],
    },
    "manualv3": {
        "template": "软件用户手册.docx",
        "md": "docs/软件用户手册v3.md",
        "out": "软件用户手册v3.docx",
        "head": {2: "1", 3: "2", 4: "3", 5: "4"},
        "body": "a4",
        "version_replaces": [("0.0.0.1", "V1.0")],
    },
    "plan": {
        "template": "系统测试说明模板.docx",
        "md": "docs/系统测试说明大纲.md",
        "out": "系统测试说明_0616v2.docx",
        "head": {2: "1", 3: "2", 4: "30", 5: "4"},  # Heading1=id1, H2=id2, H3=id30, H4=id4
        "body": "af7",  # Normal Indent
        "version_replaces": [],
    },
    "report": {
        "template": "系统测试报告模板.docx",
        "md": "docs/系统测试报告.md",
        "out": "系统测试报告_0616v2.docx",
        "head": {2: "2", 3: "4", 4: "5", 5: "6"},   # Heading1=id2, H2=id4, H3=id5, H4=id6
        "body": "3",   # Normal Indent
        "version_replaces": [],
    },
    "dagang": {
        "template": "系统测试说明模板.docx",
        "md": "docs/系统测试大纲v3.md",
        "out": "系统测试大纲v3.docx",
        "head": {2: "1", 3: "2", 4: "30", 5: "4"},  # 同 plan：说明模板 Heading1-4
        "body": "af7",  # Normal Indent
        "version_replaces": [],
    },
}

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
IMG_RE = re.compile(r"^\!\[(.*?)\]\((.+?)\)\s*$")


# ---------- XML helpers ----------
def set_pstyle(p, style_id: str) -> None:
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn("w:pStyle")):
        pPr.remove(e)
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)


def shade_para(p, fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def mono_run(p, text: str, size: int = 9) -> None:
    run = p.add_run(text)
    run.font.name = "Consolas"; run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas"); rfonts.set(qn("w:hAnsi"), "Consolas"); rfonts.set(qn("w:eastAsia"), "Consolas")


def add_rich(p, text: str) -> None:
    if not text:
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        else:
            mono_run(p, tok[1:-1])
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "808080")
        borders.append(e)
    tblPr.append(borders)


def shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# ---------- content builders ----------
def add_heading(doc, text: str, style_id: str) -> None:
    p = doc.add_paragraph(); set_pstyle(p, style_id); p.add_run(text)


def add_body(doc, text: str, body_style: str) -> None:
    p = doc.add_paragraph(); set_pstyle(p, body_style); add_rich(p, text)


def add_bullet(doc, text: str, body_style: str) -> None:
    p = doc.add_paragraph(); set_pstyle(p, body_style); add_rich(p, "• " + text)


def add_code(doc, lines: list[str], body_style: str) -> None:
    p = doc.add_paragraph(); set_pstyle(p, body_style); shade_para(p, "F5F5F5")
    for i, ln in enumerate(lines):
        mono_run(p, ln)
        if i < len(lines) - 1:
            p.add_run().add_break()


def add_md_table(doc, rows: list[list[str]]) -> None:
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header)); table_borders(t)
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; shade_cell(c, "E8EEF7")
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h); r.bold = True
    for row in body:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            if j < len(cells):
                cells[j].paragraphs[0].clear()
                add_rich(cells[j].paragraphs[0], val)


def add_image(doc, caption: str, path: Path) -> None:
    abs_path = (ROOT / path) if not path.is_absolute() else path
    if not abs_path.exists():
        p = doc.add_paragraph(); set_pstyle(p, "Normal"); p.add_run(f"[图片缺失: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = 1  # CENTER
    run = p.add_run()
    # 竖图按高度限幅，横图按宽度
    with Image.open(abs_path) as im:
        w, h = im.size
    if h > w:
        run.add_picture(str(abs_path), height=Cm(16))
    else:
        run.add_picture(str(abs_path), width=Cm(15))
    if caption:
        cap = doc.add_paragraph(); set_pstyle(cap, "Normal")
        cap.alignment = 1
        cr = cap.add_run(caption); cr.bold = True; cr.font.size = Pt(9)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", s)) and "-" in s


# ---------- main pipeline ----------
def clear_content(doc) -> None:
    """删除首个 Heading 段落及其后全部正文（保留末尾 sectPr / 承载分节的段落）。"""
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    for idx, ch in enumerate(children):
        if ch.tag.split("}")[-1] != "p":
            continue
        para = next((p for p in doc.paragraphs if p._p is ch), None)
        name = para.style.name if para is not None and para.style else ""
        if re.match(r"^Heading \d+$", name):
            start = idx
            break
    if start is None:
        raise RuntimeError("未找到正文起始 Heading，模板结构异常")
    for ch in children[start:]:
        tag = ch.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        if tag == "p":
            pPr = ch.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue
        body.remove(ch)


def fill_from_md(doc, md: Path, cfg: dict) -> None:
    lines = md.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    body_style = cfg["body"]
    head = cfg["head"]
    started = False
    while i < n:
        line = lines[i]
        s = line.strip()
        # 跳过首个 H1（封面已有的文档标题）与元信息，从首个 ## 章节开始
        if s.startswith("## ") and not started:
            started = True
        if not started:
            i += 1; continue

        # 图片
        m_img = IMG_RE.match(s)
        if m_img:
            add_image(doc, m_img.group(1), Path(m_img.group(2)))
            i += 1; continue
        # 代码块
        if s.startswith("```"):
            i += 1; block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            i += 1
            add_code(doc, block, body_style); continue
        # 标题
        m = re.match(r"^(#{2,5})\s+(.*)$", s)
        if m:
            add_heading(doc, m.group(2), head[len(m.group(1))])
            i += 1; continue
        # 表格
        if s.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                if not is_sep(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_md_table(doc, rows); continue
        # 列表
        if s.startswith("- "):
            add_bullet(doc, s[2:], body_style); i += 1; continue
        # 引用
        if s.startswith("> "):
            add_body(doc, s[2:], body_style); i += 1; continue
        if not s:
            i += 1; continue
        add_body(doc, s, body_style); i += 1


def remove_empty_headings(doc) -> None:
    """删除模板残留的空标题段落（避免被自动编号成空白"1"）。"""
    for p in list(doc.paragraphs):
        nm = p.style.name if p.style else ""
        if re.match(r"^Heading \d+$", nm) and not p.text.strip():
            p._p.getparent().remove(p._p)


def set_update_fields(doc) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
        settings.append(uf)


def apply_version(doc, replaces: list[tuple[str, str]]) -> None:
    if not replaces:
        return
    def walk(paras):
        for p in paras:
            for r in p.runs:
                if r.text:
                    for a, b in replaces:
                        r.text = r.text.replace(a, b)
    walk(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                walk(cell.paragraphs)


def main() -> int:
    key = sys.argv[1] if len(sys.argv) > 1 else ""
    if key not in DOCS:
        print(f"usage: python fill_template.py <{'|'.join(DOCS)}>", file=sys.stderr)
        return 1
    cfg = DOCS[key]
    doc = Document(cfg["template"])
    clear_content(doc)
    fill_from_md(doc, ROOT / cfg["md"], cfg)
    remove_empty_headings(doc)
    apply_version(doc, cfg["version_replaces"])
    set_update_fields(doc)
    out = ROOT / cfg["out"]
    doc.save(str(out))
    print(f"saved: {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
